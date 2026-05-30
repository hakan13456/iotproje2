#!/usr/bin/env python3
"""
Sunum rehberi — MQTT-SN vs CoAP final projesi
Sadece sana özel, nasıl anlatacağını gösterir.
"""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent.parent / "results" / "sunum_rehberi_SANA_OZEL.docx"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def baslik(doc, text, renk="1565C0"):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(
            int(renk[0:2],16), int(renk[2:4],16), int(renk[4:6],16))
    return p


def altbaslik(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return p


def yaz(doc, text, bold=False, italic=False, renk=None, boyut=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(boyut)
    if renk:
        run.font.color.rgb = RGBColor(int(renk[0:2],16), int(renk[2:4],16), int(renk[4:6],16))
    return p


def kutu(doc, baslik_text, icerik_lines, bg="E3F2FD", baslik_bg="1565C0"):
    """Renkli bilgi kutusu"""
    tbl = doc.add_table(rows=1+len(icerik_lines), cols=1)
    tbl.style = "Table Grid"
    # Başlık
    hcell = tbl.rows[0].cells[0]
    set_cell_bg(hcell, baslik_bg)
    hp = hcell.paragraphs[0]
    hr = hp.add_run(baslik_text)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    hr.font.size = Pt(11)
    # İçerik
    for i, line in enumerate(icerik_lines):
        cell = tbl.rows[i+1].cells[0]
        set_cell_bg(cell, bg)
        cp = cell.paragraphs[0]
        cr = cp.add_run(line)
        cr.font.size = Pt(10)
    doc.add_paragraph()


def madde(doc, items, renk_baslik=None):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            bold_part, normal_part = item
            r1 = p.add_run(bold_part)
            r1.bold = True
            r1.font.size = Pt(10)
            if renk_baslik:
                r1.font.color.rgb = RGBColor(
                    int(renk_baslik[0:2],16),
                    int(renk_baslik[2:4],16),
                    int(renk_baslik[4:6],16))
            r2 = p.add_run(normal_part)
            r2.font.size = Pt(10)
        else:
            r = p.add_run(item)
            r.font.size = Pt(10)


def karsilastirma_tablosu(doc):
    headers = ["Özellik", "MQTT-SN", "CoAP"]
    satirlar = [
        ["Mimari",             "Yayın/Abone (Pub/Sub)\nBroker üzerinden",     "İstek/Yanıt (Request/Response)\nDoğrudan iletişim"],
        ["Broker gerekir mi?", "EVET — merkezi broker şart",                  "HAYIR — broker yok"],
        ["Paket onayı (ACK)",  "QoS-0: yok\nQoS-1: PUBACK onayı var",         "CON: ACK alınana kadar tekrar dener\nNON: onaysız"],
        ["Başlık boyutu",      "2–5 bayt (çok küçük)",                         "4+ bayt"],
        ["Topoloji uyumu",     "Ağaç topolojisinde iyi\n(broker toplar)",      "Yıldız topolojisinde mükemmel"],
        ["Bizim ağaç sonucu",  "PDR: %100\nRTT: 15.9 ms",                      "PDR: %84.8\nRTT: 21.8 ms"],
        ["Bizim yıldız sonucu","PDR: %38.5 (PINGRESP)\nBroker aşırı yüklendi","PDR: %100\nRTT: 17.7 ms"],
        ["Kısaca",             "Merkezi, verimli, broker bağımlı",             "Dağıtık, esnek, broker yok"],
    ]
    col_w = [Cm(3.8), Cm(6.5), Cm(6.5)]
    tbl = doc.add_table(rows=1+len(satirlar), cols=3)
    tbl.style = "Table Grid"
    # Header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.width = col_w[i]
        set_cell_bg(cell, "1565C0")
        r = cell.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Satırlar
    for ri, row in enumerate(satirlar):
        bg = "EEF2FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            cell.width = col_w[ci]
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            if ci == 0:
                r.font.bold = True
    doc.add_paragraph()


def build():
    doc = Document()
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ══════════════════════════════════════════════════
    # KAPAK
    # ══════════════════════════════════════════════════
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SUNUM REHBERİ — SADECE SANA ÖZEL")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0D,0x47,0xA1)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("MQTT-SN vs CoAP — FIT IoT-LAB Final Projesi\nNasıl Anlatacaksın, Ne Diyeceksin, Terimler Ne Demek")
    r2.font.size = Pt(13)
    r2.font.bold = True

    doc.add_paragraph()
    yaz(doc, "Bu dosya sana özel yazıldı. Hocaya veya sunuma gitmez. "
             "Projeyi nasıl anlatacağını, sorulan sorulara ne cevap vereceğini ve "
             "teknik terimlerin ne anlama geldiğini adım adım açıklıyor.",
        italic=True, renk="555555")
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 1. PROJE NE HAKKINDA — 3 CÜMLEDE ANLAT
    # ══════════════════════════════════════════════════
    baslik(doc, "1. Projeyi 3 Cümlede Nasıl Anlat")

    kutu(doc, "Sunuma başlarken bunu söyle:", [
        '"Bu projede iki farklı IoT iletişim protokolünü — MQTT-SN ve CoAP\'ı —',
        ' gerçek donanım üzerinde karşılaştırdım."',
        '"FIT IoT-LAB adlı Fransız araştırma altyapısında 15 adet kablosuz sensör düğümü kullandım."',
        '"Üç farklı ağ şekli (ağaç, yıldız, doğrusal) kurdum ve her protokolün ne kadar hızlı,',
        ' ne kadar güvenilir çalıştığını ölçtüm."',
    ])

    yaz(doc, "Soru gelirsе: 'Neden bu iki protokol?'", bold=True)
    yaz(doc, "→ Çünkü ikisi de IoT dünyasında en çok kullanılan hafif protokoller. "
             "Biri broker'a ihtiyaç duyuyor (MQTT-SN), diğeri duymуyor (CoAP). "
             "Bu fark performansı nasıl etkiliyor diye baktık.")
    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 2. ANAHTAR TERİMLER
    # ══════════════════════════════════════════════════
    baslik(doc, "2. Anahtar Terimler — Ne Demek, Nasıl Açıklarsın")

    terimler = [
        ("IoT (Nesnelerin İnterneti)", "E3F2FD",
         "Fiziksel cihazların internete bağlanması. Termometre, sensör, akıllı ampul — "
         "bunların hepsi IoT. Biz 15 kablosuz sensör düğümü kullandık.",
         "Soru: 'IoT nedir?' → 'Günlük cihazların internete bağlanarak veri toplaması veya kontrol edilmesi.'"),

        ("FIT IoT-LAB", "E8F5E9",
         "Fransa'da bulunan devlet destekli araştırma altyapısı. "
         "Yüzlerce gerçek donanım düğümü var, uzaktan SSH ile bağlanıp deney yapıyorsun. "
         "Biz Saclay sitesindeki A8-M3 düğümlerini kullandık.",
         "Soru: 'Donanımı nereden aldın?' → 'Kendi bilgisayarımdan FIT IoT-LAB\'a SSH ile bağlandım, "
         "oradaki gerçek fiziksel düğümleri uzaktan kontrol ettim.'"),

        ("A8-M3 Düğümü", "FFF8E1",
         "İki işlemcili özel bir kart. İçinde iki ayrı işlemci var:\n"
         "  - ARM Cortex-A8: Linux çalıştırıyor, bizim broker ve ağ geçidimiz buraya kuruldu\n"
         "  - ARM Cortex-M3: RIOT OS çalıştırıyor, asıl sensör kodu buraya yüklendi\n"
         "AT86RF231 adlı 802.15.4 radyo çipi var, WiFi değil.",
         "Soru: 'Hangi donanım?' → 'FIT IoT-LAB A8-M3 düğümleri. "
         "Her birinde iki işlemci var: Linux tarafı ve RIOT OS tarafı.'"),

        ("RIOT OS", "FCE4EC",
         "Kısıtlı (az RAM/CPU) cihazlar için yazılmış açık kaynak işletim sistemi. "
         "Arduino gibi ama çok daha güçlü — gerçek bir OS. "
         "6LoWPAN, RPL, CoAP, MQTT-SN modülleri hazır geliyor.",
         "Soru: 'Neden RIOT OS?' → 'Kısıtlı IoT donanımları için optimize edilmiş, "
         "CoAP ve MQTT-SN kütüphaneleri built-in geliyor.'"),

        ("IEEE 802.15.4 / 6LoWPAN", "F3E5F5",
         "802.15.4: Kablosuz sensör ağları için fiziksel radyo standardı. WiFi gibi ama çok daha az güç tüketiyor.\n"
         "6LoWPAN: IPv6 paketlerini bu küçük 802.15.4 çerçevelerine sığdırmak için sıkıştırma katmanı.",
         "Soru: 'WiFi neden kullanmadın?' → '802.15.4 sensörler için özel, çok daha az güç tüketiyor. "
         "6LoWPAN ile IPv6 adresler doğrudan sensörde çalışıyor.'"),

        ("RPL (Routing Protocol)", "E8EAF6",
         "Düşük güçlü ağlar için IPv6 yönlendirme protokolü (RFC 6550). "
         "Otomatik olarak ağaç yapısı (DODAG) oluşturuyor. "
         "Border router kök düğüm, diğerleri altına hiyerarşik bağlanıyor.",
         "Soru: 'Ağaç topolojiyi nasıl kurдun?' → 'RPL protokolü otomatik oluşturdu. "
         "TX güç seviyelerini düşürerek katmanları zorladım.'"),

        ("PDR — Paket İletim Oranı", "E0F7FA",
         "Gönderilen paket sayısına karşı başarıyla alınan paket sayısının oranı.\n"
         "PDR = Alınan / Gönderilen × 100\n"
         "PDR %100 = hiç kayıp yok. PDR %61 = 100 paketten 61'i ulaşmış.",
         "Soru: 'PDR ne demek?' → 'Packet Delivery Ratio. "
         "Gönderilen paketin hedefe ulaşma yüzdesi. %100 mükemmel, düşükse kayıp var.'"),

        ("RTT — Gidiş-Dönüş Gecikmesi", "FFF3E0",
         "Round Trip Time. Paket gönderilip yanıt alınana kadar geçen toplam süre (ms cinsinden).\n"
         "RTT = Yanıt alındı zamanı − Paket gönderildi zamanı\n"
         "Düşük RTT = hızlı. Yüksek RTT = gecikmeli.",
         "Soru: 'RTT ne?' → 'Round Trip Time. Pingdeki gibi — paketi gönderip yanıt gelene kadar geçen süre.'"),
    ]

    for terim, bg, aciklama, soru_cevap in terimler:
        altbaslik(doc, terim)
        kutu(doc, "Ne demek?", aciklama.split('\n'), bg=bg)
        p = doc.add_paragraph()
        r1 = p.add_run("Hoca sorarsa: ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
        r1.font.size = Pt(10)
        r2 = p.add_run(soru_cevap)
        r2.font.size = Pt(10)
        r2.italic = True
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 3. MQTT-SN vs CoAP FARKI
    # ══════════════════════════════════════════════════
    baslik(doc, "3. MQTT-SN ile CoAP Arasındaki Fark — En Önemli Kısım")

    yaz(doc, "Bu soruyu kesinlikle sorarlar. Şöyle anlat:", bold=True, renk="C62828")

    kutu(doc, "MQTT-SN nedir? (Bunu söyle)", [
        "MQTT-SN = MQTT for Sensor Networks.",
        "Normal MQTT'nin kablosuz sensörler için hafifletilmiş versiyonu.",
        "Çalışma şekli: Sensör → Broker'a mesaj gönderir → Broker → Abone olan herkese dağıtır.",
        "Ortada mutlaka bir BROKER (aracı sunucu) olması lazım.",
        "Bizde broker = RSMB (Really Small Message Broker), border router'ın Linux tarafında çalıştı.",
        "QoS seviyeleri var: QoS-0 (onaysız, hızlı), QoS-1 (PUBACK ile onaylı, güvenilir).",
    ])

    kutu(doc, "CoAP nedir? (Bunu söyle)", [
        "CoAP = Constrained Application Protocol.",
        "HTTP'nin kısıtlı cihazlar için UDP tabanlı versiyonu.",
        "Çalışma şekli: İstemci → Sunucuya doğrudan istek → Sunucu → Yanıt.",
        "BROKER GEREKMİYOR. Düğümler birbirleriyle doğrudan konuşuyor.",
        "CON (Confirmable) mesaj = ACK alınana kadar tekrar dener → güvenilir.",
        "NON (Non-confirmable) = onaysız, hızlı ama kayıp olabilir.",
        "Bizde: 1 düğüm sunucu, 13 düğüm istemci olarak çalıştı.",
    ], bg="FFF8E1")

    doc.add_paragraph()
    yaz(doc, "Karşılaştırma Tablosu:", bold=True)
    karsilastirma_tablosu(doc)

    kutu(doc, "Bunu da ekle — gerçek hayatta hangisi kullanılır?", [
        "MQTT-SN: Akıllı fabrika, merkezi veri toplama, birçok sensörden tek noktaya.",
        "CoAP: Akıllı ev, REST API benzeri erişim, her cihaza doğrudan ulaşmak istediğinde.",
        "İkisi de IETF/OASIS standartları — endüstride gerçekten kullanılıyor.",
    ], bg="E8F5E9")

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 4. TOPOLOJILER NE DEMEK
    # ══════════════════════════════════════════════════
    baslik(doc, "4. Topolojiler — Neden 3 Farklı Şekil Test Ettik")

    yaz(doc, "Topoloji = Düğümlerin birbirine bağlanma şekli. Ağın 'haritası'.", bold=True)
    doc.add_paragraph()

    kutu(doc, "Ağaç Topolojisi (Tree) — En Önemli", [
        "RPL protokolü ile otomatik oluştu.",
        "Border router = kök (root). Diğerleri dallar ve yapraklar.",
        "Katmanlar: Kök → Kademe1 → Kademe2 → Kademe3 (yapraklar)",
        "Paket çoklu atlama yapıyor: yaprak → kademe2 → kademe1 → kök",
        "SONUÇ: MQTT-SN burada çok iyi çalıştı (%100 PDR, 15.9ms)",
        "Çünkü broker tüm trafiği topluyor, ağaç yapısıyla uyumlu.",
    ])

    kutu(doc, "Yıldız Topolojisi (Star)", [
        "Tüm sensörler doğrudan merkeze (border router) bağlı.",
        "Tek atlama — en basit yapı.",
        "SONUÇ: CoAP burada kazandı (%100 PDR, 17.7ms)",
        "MQTT-SN'de broker aynı anda çok bağlantıyı kaldıramadı (%38.5 PDR).",
    ], bg="FFF8E1")

    kutu(doc, "Doğrusal Topoloji (Linear)", [
        "Düğümler zincir gibi sıralı: A → B → C → D → ...",
        "Paketin çok atlama yapması gerekiyor.",
        "SONUÇ: Her iki protokolde de PDR düştü.",
        "CoAP %61.9 PDR — çok atlamalı kayıplar.",
        "MQTT-SN QoS-0 kullandı, ölçüm sınırlı kaldı.",
    ], bg="FCE4EC")

    doc.add_paragraph()
    yaz(doc, "Neden 3 topoloji? → Gerçek dünya ağları farklı şekillerde kurulabiliyor. "
             "Hangi protokol hangi yapıda daha iyi çalışıyor diye test ettik.", italic=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 5. WIRESHARK / PCAP NE YAPTIn
    # ══════════════════════════════════════════════════
    baslik(doc, "5. Wireshark / Paket Analizi — Ne Yaptın, Nasıl Anlat")

    yaz(doc, "Soru: 'Wireshark ile ne yaptın?'", bold=True, renk="C62828")

    kutu(doc, "Şöyle anlat:", [
        "A8 düğümlerinin Linux tarafında 'tshark' aracı var (Wireshark'ın terminal versiyonu).",
        "wpan0 arayüzünü dinledim → 802.15.4 kablosuz trafiğini yakaladım.",
        "Yakalanan trafik .pcap dosyası olarak kaydedildi.",
        "Bu .pcap dosyasını Python Scapy kütüphanesiyle analiz ettim:",
        "  → CoAP mesajlarını Message-ID'ye göre eşleştirdim → RTT hesapladım",
        "  → MQTT-SN PUBLISH/PUBACK çiftlerini eşleştirdim → RTT ve PDR hesapladım",
        "Sonuçları grafiğe döktüm.",
    ])

    kutu(doc, "Pcap dosyalarında ne gördün? (Somut örnek ver)", [
        "mqttsn_tree.pcap: 562 paket, 144.5 saniye, PUBLISH:11, PUBACK:11 → %100 PDR",
        "coap_tree.pcap: 163 paket, CON:33, ACK:28 → %84.8 PDR",
        "mqttsn_star.pcap: 8 kez yeniden bağlanma (CONNECT/CONNACK) → broker aşırı yüklü",
        "coap_star.pcap: 31 CON, 31 ACK → mükemmel, hiç kayıp yok",
    ], bg="E8F5E9")

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 6. SONUÇLAR — NASIL YORUMLA
    # ══════════════════════════════════════════════════
    baslik(doc, "6. Sonuçları Nasıl Yorumlarsın")

    yaz(doc, "Sonuç sorusu gelirse bunu söyle:", bold=True)

    kutu(doc, "Sunum sonu sonuç cümlesi:", [
        '"Her iki protokol de kısıtlı IoT ortamlarında çalışıyor,',
        ' ama farklı topolojilerde üstünlük değişiyor."',
        '"MQTT-SN, broker sayesinde ağaç topolojisinde daha güvenilir:"',
        '  Ağaç: MQTT-SN %100 PDR, CoAP %84.8 PDR',
        '"CoAP ise broker gerektirmediğinden yıldız topolojisinde daha iyi:"',
        '  Yıldız: CoAP %100 PDR, MQTT-SN %38.5 PDR',
        '"Protokol seçimi ağ yapısına göre değişmeli."',
    ])

    yaz(doc, "Soru: 'MQTT-SN neden yıldızda kötü çalıştı?'", bold=True, renk="C62828")
    yaz(doc, "→ 15 düğüm aynı anda broker'a bağlanmaya çalıştı. "
             "Pcap'te 8 kez CONNECT/CONNACK gördüm — broker sürekli bağlantı kesip yeniden kurdu. "
             "Bu broker'ın aşırı yüklendiğini gösteriyor.")

    yaz(doc, "Soru: 'CoAP neden ağaçta daha kötü?'", bold=True, renk="C62828")
    yaz(doc, "→ CoAP'ta her istek doğrudan sunucuya gitmek zorunda. "
             "Çok atlamalı ağaç yapısında her atlama bir kayıp riski. "
             "5 CON paketi ACK alamadan timeout oldu.")

    yaz(doc, "Soru: 'QoS-1 ve QoS-0 nedir?'", bold=True, renk="C62828")
    yaz(doc, "→ QoS-0: paket gönder, onay bekleme — hızlı ama kayıp olabilir. "
             "QoS-1: paket gönder, PUBACK (onay) gel — yavaş ama güvenilir. "
             "Ağaç deneyimizde QoS-1 kullandık, bu yüzden %100 PDR aldık.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 7. SUNUM AKIŞI — ADIM ADIM
    # ══════════════════════════════════════════════════
    baslik(doc, "7. Sunumu Adım Adım Nasıl Yap")

    adimlar = [
        ("1. Açılış (30 saniye)", [
            "'Bu projede iki IoT protokolünü — MQTT-SN ve CoAP — gerçek donanımda karşılaştırdım.'",
            "'FIT IoT-LAB'da 15 A8-M3 düğümü kullandım, RIOT OS çalıştırdım.'",
        ]),
        ("2. Neden bu protokoller? (1 dakika)", [
            "'MQTT-SN ve CoAP, kısıtlı sensörler için en yaygın iki protokol.'",
            "'Biri broker istiyor (MQTT-SN), biri istemiyor (CoAP). Bu farkın performansı nasıl etkilediğini ölçtük.'",
        ]),
        ("3. Kurulum (1-2 dakika)", [
            "'3 topoloji kurdum: ağaç, yıldız, doğrusal.'",
            "'Ağaçta RPL otomatik hiyerarşi oluşturdu, TX güç ile 3 kademe zorladım.'",
            "'MQTT-SN için border router'da RSMB broker çalıştırdım.'",
        ]),
        ("4. Ölçümler (2 dakika)", [
            "'RTT: paketi gönder, geri gel — farkı millisaniye cinsinden ölçtüm.'",
            "'PDR: gönderilen/alınan paket oranı.'",
            "'Wireshark (tshark) ile 802.15.4 trafiğini yakaladım, Python Scapy ile analiz ettim.'",
        ]),
        ("5. Sonuçlar (2 dakika)", [
            "Tabloyu göster: 'Ağaçta MQTT-SN kazandı: %100 PDR, 15.9ms'",
            "'Yıldızda CoAP kazandı: %100 PDR, 17.7ms'",
            "'Topoloji seçimi protokol performansını doğrudan etkiliyor.'",
        ]),
        ("6. Sonuç (30 saniye)", [
            "'MQTT-SN merkezi ağaç yapılarında, CoAP dağıtık yıldız yapılarında daha uygun.'",
            "'Gerçek dünyada seçim: ağ mimarisine ve QoS gereksinimlerine göre yapılmalı.'",
        ]),
    ]

    for baslik_text, maddeler in adimlar:
        kutu(doc, baslik_text, maddeler, bg="E8F5E9", baslik_bg="2E7D32")

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 8. TAHMİNİ SORULAR VE CEVAPLAR
    # ══════════════════════════════════════════════════
    baslik(doc, "8. Hocadan Gelebilecek Sorular ve Cevaplar")

    sorular = [
        ("FIT IoT-LAB nedir, neden oraya giriş izni var?",
         "Fransa'da INRIA tarafından kurulan açık erişimli araştırma altyapısı. "
         "Ücretsiz hesap açılıyor, akademik projeler için erişim veriyorlar. "
         "Saclay'da 400'den fazla A8-M3 düğümü var."),

        ("Gerçek deney mi yaptın yoksa simülasyon mu?",
         "Gerçek deney. Fiziksel A8-M3 donanımlarına SSH ile bağlandım, "
         "ELF dosyasını flash'ladım, 802.15.4 radyo üzerinden gerçek paket trafiği oluştu. "
         "Simülasyon değil."),

        ("Neden sadece 11 MQTT-SN paketi var ağaçta?",
         "Deney süresi kısıtlıydı ve QoS-1 yavaş çalışıyor — broker onay bekliyor. "
         "Önemli olan 11 paketin hepsinin ulaşmış olması: %100 PDR."),

        ("CoAP'ta 5 paket neden kayboldu?",
         "3 kademeli ağaçta her atlama kayıp riski. 5 CON paketi 2 saniyelik timeout içinde "
         "ACK alamadı. RPL ağacının kararlı olmadığı anlarda düğüm bağlantısı kesilebiliyor."),

        ("emcute ve gcoap ne?",
         "RIOT OS'un built-in kütüphaneleri. emcute = MQTT-SN client, gcoap = CoAP client/server. "
         "Tutorial örneklerini aldım, üzerine timestamp ve CSV çıktı ekledim."),

        ("Hangi dilde kod yazdın?",
         "Firmware: C (RIOT OS API). Otomasyon scriptleri: Bash. "
         "Veri analizi ve rapor üretimi: Python (Scapy, matplotlib, ReportLab, python-docx)."),

        ("QoS-0 ile QoS-1 arasında ne fark var, neden ağaçta QoS-1 seçtin?",
         "QoS-0: onaysız, hızlı ama kayıp olabilir. "
         "QoS-1: PUBACK ile onaylı, PDR güvencesi var ama daha yavaş. "
         "Ağaçta güvenilirlik ölçmek istediğim için QoS-1 seçtim."),

        ("Topoloji farkı neden önemli?",
         "Gerçek ortamlarda ağlar farklı şekillerde kurulabiliyor. "
         "Fabrika = ağaç, akıllı ev = yıldız, dağıtık sensör hattı = doğrusal. "
         "Hangi topolojide hangi protokol daha iyi diye baktık."),
    ]

    for soru, cevap in sorular:
        p = doc.add_paragraph()
        r1 = p.add_run("S: " + soru)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        p2 = doc.add_paragraph()
        r2 = p2.add_run("C: " + cevap)
        r2.font.size = Pt(10)
        r2.italic = True
        p2.paragraph_format.left_indent = Cm(0.8)
        p2.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 9. HIZLI ÖZET — SON KONTROL
    # ══════════════════════════════════════════════════
    baslik(doc, "9. Hızlı Özet — Sunuma 5 Dakika Kala Bunu Oku")

    kutu(doc, "HATIRLA:", [
        "MQTT-SN = pub/sub, broker gerekli, QoS ile güvenilir iletim",
        "CoAP = request/response, broker yok, CON/ACK ile güvenilir iletim",
        "Ağaç topolojisi: MQTT-SN kazandı (%100 PDR, 15.9ms RTT)",
        "Yıldız topolojisi: CoAP kazandı (%100 PDR, 17.7ms RTT)",
        "PDR = ulaşan paket yüzdesi | RTT = gidiş-dönüş süresi ms",
        "FIT IoT-LAB = gerçek Fransız araştırma donanımı, SSH ile bağlandım",
        "RIOT OS = kısıtlı IoT cihazları için OS | A8-M3 = çift işlemcili düğüm",
        "Wireshark (tshark) ile 802.15.4 trafiği yakaladım, Python Scapy ile analiz ettim",
    ], bg="FFEB3B", baslik_bg="F57F17")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Başarılar! Bu projeyi sen yaptın, anlat.")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    doc.save(str(OUT))
    print(f"Sunum rehberi olusturuldu -> {OUT}")


if __name__ == "__main__":
    build()
