#!/usr/bin/env python3
"""
Word (.docx) deney raporu — MQTT-SN vs CoAP, FIT IoT-LAB
"""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RESULTS = Path(__file__).parent.parent / "results"
OUT     = RESULTS / "deney_raporu_mqttsn_vs_coap.docx"

DATA = [
    ("MQTT-SN", "Ağaç",    562,  11, 11, 100.0, 15.9, 15.6, 19.6, "QoS-1, PUBACK"),
    ("MQTT-SN", "Doğrusal",726,   2,  2, 100.0,  1.8,  1.8,  1.8, "QoS-0, PINGRESP"),
    ("MQTT-SN", "Yıldız",  680,  13,  5,  38.5,  2.0,  2.0,  2.1, "QoS-0, PINGRESP"),
    ("CoAP",    "Ağaç",    163,  33, 28,  84.8, 21.8, 20.5, 31.1, "CON/ACK"),
    ("CoAP",    "Doğrusal",147,  21, 13,  61.9, 18.0, 17.7, 19.7, "CON/ACK"),
    ("CoAP",    "Yıldız",   79,  31, 31, 100.0, 17.7, 17.4, 19.0, "CON/ACK"),
]

HEADER_RGB = RGBColor(0x15, 0x65, 0xC0)  # dark blue


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return h


def add_paragraph(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_results_table(doc):
    headers = ["Protokol","Topoloji","Toplam\nPaket","Gönderilen",
               "Alınan","PDR (%)","RTT Ort.\n(ms)","RTT Med.\n(ms)","RTT p95\n(ms)","Not"]
    col_w   = [Cm(2.2), Cm(2.2), Cm(1.6), Cm(1.8), Cm(1.5),
               Cm(1.5), Cm(1.8), Cm(1.8), Cm(1.8), Cm(3.5)]
    t = doc.add_table(rows=1+len(DATA), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, (h, w) in enumerate(zip(headers, col_w)):
        cell = t.rows[0].cells[i]
        cell.width = w
        cell.text = h
        set_cell_bg(cell, "1565C0")
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row_data in enumerate(DATA):
        row = t.rows[ri+1]
        bg = "EEF2FF" if ri % 2 == 0 else "FFFFFF"
        vals = [row_data[0], row_data[1], str(row_data[2]),
                str(row_data[3]), str(row_data[4]),
                f"{row_data[5]:.1f}", f"{row_data[6]:.1f}",
                f"{row_data[7]:.1f}", f"{row_data[8]:.1f}", row_data[9]]
        for ci, (val, w) in enumerate(zip(vals, col_w)):
            cell = row.cells[ci]
            cell.width = w
            cell.text = val
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return t


def add_image(doc, path, caption, width_cm=15):
    if Path(path).exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(8)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ── Kapak ─────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_heading("DENEY RAPORU", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("MQTT-SN ve CoAP Protokollerinin FIT IoT-LAB Üzerinde Karşılaştırmalı Analizi")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].bold = True

    doc.add_paragraph()
    info = doc.add_paragraph(
        f"Öğrenci: Hakan Dönmez   |   No: 20010011038\n"
        f"Tarih: {date.today().strftime('%d.%m.%Y')}\n"
        "FIT IoT-LAB — Saclay Sitesi, 15× A8-M3 Düğümü\n"
        "RIOT OS | 6LoWPAN/802.15.4 | RPL Routing"
    )
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── 1. Giriş ─────────────────────────────────────────────────
    add_heading(doc, "1. Giriş")
    add_paragraph(doc,
        "Bu deney raporu, kısıtlı IoT ortamlarında yaygın olarak kullanılan iki uygulama "
        "katmanı protokolünün — MQTT-SN ve CoAP — performansını FIT IoT-LAB testbed üzerinde "
        "karşılaştırmalı olarak değerlendirmektedir. Deneyler üç farklı ağ topolojisinde "
        "(ağaç, yıldız, doğrusal) 15 adet A8-M3 düğümü ile gerçekleştirilmiştir. "
        "Temel performans metrikleri olarak Paket İletim Oranı (PDR), Gidiş-Dönüş Gecikmesi (RTT) "
        "ve paket yapısı Wireshark/tshark ile analiz edilmiştir.")

    # ── 2. Deneysel Kurulum ───────────────────────────────────────
    add_heading(doc, "2. Deneysel Kurulum")
    add_heading(doc, "2.1 Donanım ve Platform", level=2)
    add_paragraph(doc,
        "FIT IoT-LAB (Saclay sitesi) üzerinde çalışan 15× IoT-LAB A8-M3 düğümü kullanılmıştır. "
        "Her düğüm ARM Cortex-A8 (Linux) + ARM Cortex-M3 (RIOT OS) + AT86RF231 802.15.4 radyosundan "
        "oluşmaktadır. RF kanalı 26 (2.480 GHz), yönlendirme protokolü RPL (RFC 6550)'dir.")

    setup = [
        ["Platform","FIT IoT-LAB Saclay"],
        ["Düğüm tipi","IoT-LAB A8-M3 (Cortex-A8 + Cortex-M3)"],
        ["İşletim sistemi","RIOT OS"],
        ["Radyo","IEEE 802.15.4 — AT86RF231, Kanal 26"],
        ["MQTT-SN kütüphanesi","emcute (RIOT built-in)"],
        ["CoAP kütüphanesi","gcoap (RIOT built-in)"],
        ["MQTT-SN broker","mosquitto RSMB (A8 Linux)"],
        ["Border router","a8-103"],
        ["Sensör düğümleri","14× (a8-97 .. a8-112)"],
    ]
    tbl = doc.add_table(rows=len(setup), cols=2)
    tbl.style = "Table Grid"
    for ri, (k, v) in enumerate(setup):
        tbl.rows[ri].cells[0].text = k
        tbl.rows[ri].cells[1].text = v
        set_cell_bg(tbl.rows[ri].cells[0], "E3F2FD")
        tbl.rows[ri].cells[0].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

    add_heading(doc, "2.2 Topoloji Tasarımı", level=2)
    add_paragraph(doc,
        "Ağaç topolojisi RPL protokolü ile otomatik oluşturulmuştur. TX güç seviyeleri ile "
        "3 kademeli hiyerarşi zorlanmıştır: Kademe 0 (kök/border router), Kademe 1 (+0 dBm), "
        "Kademe 2 (−10 dBm), Kademe 3 yaprak düğümler (−17 dBm). Yıldız topolojisinde tüm "
        "sensörler doğrudan border router'a bağlıdır. Doğrusal topolojide düğümler zincir biçimindedir.")

    # ── 3. Metodoloji ─────────────────────────────────────────────
    add_heading(doc, "3. Metodoloji")
    add_heading(doc, "3.1 MQTT-SN Deneyi", level=2)
    add_paragraph(doc,
        "Border router A8 Linux tarafında RSMB broker başlatılmıştır. M3 sensörler RIOT emcute "
        "kütüphanesi ile broker'a bağlanmış; ağaç topolojisinde QoS-1 (PUBACK onaylı), "
        "yıldız/doğrusal topolojilerde QoS-0 kullanılmıştır. RTT ölçümü QoS-1'de "
        "PUBLISH→PUBACK zaman farkı, QoS-0'da PINGREQ→PINGRESP zaman farkı ile yapılmıştır.")

    add_heading(doc, "3.2 CoAP Deneyi", level=2)
    add_paragraph(doc,
        "Bir düğüm CoAP sunucusu (gcoap), diğerleri istemci olarak yapılandırılmıştır. "
        "İstemciler /metrics/rtt kaynağına Confirmable (CON) POST göndermiştir. "
        "RTT = ACK alım zamanı − CON gönderim zamanı. PDR = ACK alınan / CON gönderilen.")

    add_heading(doc, "3.3 Wireshark / Tshark Analizi", level=2)
    add_paragraph(doc,
        "A8 Linux tarafından wpan0 arayüzü üzerinden 802.15.4 trafiği tshark ile yakalanmıştır. "
        "Yakalama dosyaları (.pcap) Python Scapy kütüphanesi ile analiz edilmiş; CoAP mesaj "
        "kimliklerine ve MQTT-SN mesaj tiplerine göre RTT ve PDR hesaplanmıştır.")

    # ── 4. Bulgular ────────────────────────────────────────────────
    add_heading(doc, "4. Bulgular")
    add_heading(doc, "4.1 Ölçüm Sonuçları", level=2)
    add_results_table(doc)
    doc.add_paragraph("Tablo 1: Tüm deneyler için PDR ve RTT ölçüm sonuçları").runs[0].font.italic = True

    # Comparison chart
    deep_png = RESULTS / "deep_comparison.png"
    if deep_png.exists():
        doc.add_paragraph()
        add_image(doc, deep_png, "Şekil 1: RTT, PDR ve paket sayısı karşılaştırması — tüm topolojiler")

    # ── Wireshark görüntüleri ──────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "4.2 Wireshark Paket Analizi Görüntüleri")
    add_paragraph(doc,
        "Aşağıda her topoloji ve protokol için 802.15.4 arayüzünden yakalanan "
        "pcap dosyalarının paket boyutu/zaman ve paketler-arası süre (IAT) grafikleri verilmiştir.")

    ws_files = [
        ("wireshark_mqttsn_tree.png",   "Şekil 2: MQTT-SN — Ağaç Topolojisi"),
        ("wireshark_mqttsn_star.png",   "Şekil 3: MQTT-SN — Yıldız Topolojisi"),
        ("wireshark_mqttsn_linear.png", "Şekil 4: MQTT-SN — Doğrusal Topoloji"),
        ("wireshark_coap_tree.png",     "Şekil 5: CoAP — Ağaç Topolojisi"),
        ("wireshark_coap_star.png",     "Şekil 6: CoAP — Yıldız Topolojisi"),
        ("wireshark_coap_linear.png",   "Şekil 7: CoAP — Doğrusal Topoloji"),
    ]
    for fname, caption in ws_files:
        path = RESULTS / fname
        add_image(doc, path, caption, width_cm=14)
        doc.add_paragraph()

    # ── 5. Analiz ──────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "5. Analiz ve Tartışma")

    add_heading(doc, "5.1 Ağaç Topolojisi", level=2)
    add_paragraph(doc,
        "MQTT-SN, QoS-1 mekanizması sayesinde %100 PDR ve 15.9 ms ortalama RTT elde etmiştir. "
        "CoAP %84.8 PDR ve 21.8 ms RTT kaydetmiştir. MQTT-SN broker üzerinden merkezi "
        "yönlendirmesi çok atlamalı ağaç yapısında daha kararlı çalışmıştır.")

    add_heading(doc, "5.2 Yıldız Topolojisi", level=2)
    add_paragraph(doc,
        "CoAP %100 PDR ve 17.7 ms RTT ile mükemmel performans gösterirken, MQTT-SN "
        "PINGRESP PDR'si %38.5'e düşmüştür. Çok düğümlü eş zamanlı bağlanma (8× CONNECT/CONNACK) "
        "broker aşırı yüklemesine yol açmış olabilir.")

    add_heading(doc, "5.3 Doğrusal Topoloji", level=2)
    add_paragraph(doc,
        "CoAP %61.9 PDR ve 18.0 ms RTT elde etmiştir. MQTT-SN QoS-0 ile çalışmış; "
        "yalnızca 2 PINGREQ/PINGRESP ölçülebildiğinden istatistiksel güvenilirlik düşüktür. "
        "Çok atlama kayıpları CoAP CON/ACK oranını azaltmaktadır.")

    add_heading(doc, "5.4 Protokol Karşılaştırması", level=2)
    add_paragraph(doc,
        "MQTT-SN: Broker tabanlı pub/sub model. QoS-1 güvenilir iletim. "
        "Küçük başlık boyutu (2-5 bayt) kısıtlı düğümler için verimli. "
        "Ağaç topolojisinde üstün.\n\n"
        "CoAP: İstek/yanıt modeli, broker gerektirmez. CON/ACK yerleşik güvenilirlik. "
        "Yıldız topolojisinde %100 PDR. Çok atlamalı senaryolarda kayıp artar.")

    # ── 6. Sonuç ──────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "6. Sonuç")
    add_paragraph(doc,
        "Bu çalışmada MQTT-SN ve CoAP FIT IoT-LAB testbedinde üç farklı topoloji altında "
        "deneysel olarak karşılaştırılmıştır:\n\n"
        "• Ağaç topolojisi: MQTT-SN üstün PDR (%100 vs %84.8) ve düşük RTT (15.9 ms vs 21.8 ms).\n"
        "• Yıldız topolojisi: CoAP %100 PDR ve 17.7 ms RTT ile daha iyi.\n"
        "• Doğrusal topoloji: CoAP %61.9 PDR ile ölçülebilir sonuç; MQTT-SN QoS-0 sınırlı veri.\n\n"
        "MQTT-SN hiyerarşik ağaç yapılarında avantajlıdır. CoAP broker gerektirmeyen, "
        "yıldız veya küçük ölçekli ağlarda tercih edilmelidir.")

    # ── 7. Kaynaklar ──────────────────────────────────────────────
    add_heading(doc, "7. Kaynaklar")
    refs = [
        '[1] Palmese et al., "CoAP vs. MQTT-SN Comparison and Performance Evaluation in '
        'Publish-Subscribe Environments," WF-IoT 2019.',
        '[2] Adjih et al., "FIT IoT-LAB: A Large Scale Open Experimental IoT Testbed," '
        'IEEE WF-IoT 2015.',
        '[3] RIOT OS Documentation — emcute / gcoap modules. https://riot-os.org',
        '[4] RFC 7252 — The Constrained Application Protocol (CoAP), IETF 2014.',
        '[5] OASIS MQTT-SN v1.2 Specification, 2013.',
        '[6] RFC 6550 — RPL: IPv6 Routing Protocol for Low-Power and Lossy Networks, 2012.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style="List Bullet")
        p.runs[0].font.size = Pt(9)

    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Bu rapor FIT IoT-LAB Saclay sitesindeki deneyler ve pcap verileri temel alınarak "
        f"üretilmiştir. Tarih: {date.today().strftime('%d.%m.%Y')}")
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.italic = True
    footer_p.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(OUT))
    print(f"Word raporu olusturuldu -> {OUT}")


if __name__ == "__main__":
    build()
