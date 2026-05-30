#!/usr/bin/env python3
"""
Optimizasyon rehberi — MQTT-SN vs CoAP projesi
Hocaya anlatmak için, deney bulgularından çıkan somut öneriler.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent.parent / "results" / "optimizasyon_onerileri.docx"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def baslik(doc, text, level=1, renk="0D47A1"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(
            int(renk[0:2],16), int(renk[2:4],16), int(renk[4:6],16))
    return h


def yaz(doc, text, bold=False, italic=False, boyut=10, renk=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(boyut)
    if renk:
        run.font.color.rgb = RGBColor(
            int(renk[0:2],16), int(renk[2:4],16), int(renk[4:6],16))
    return p


def kutu(doc, baslik_text, satirlar, bg="E3F2FD", baslik_bg="1565C0"):
    tbl = doc.add_table(rows=1 + len(satirlar), cols=1)
    tbl.style = "Table Grid"
    hcell = tbl.rows[0].cells[0]
    set_cell_bg(hcell, baslik_bg)
    hr = hcell.paragraphs[0].add_run(baslik_text)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    hr.font.size = Pt(10)
    for i, s in enumerate(satirlar):
        cell = tbl.rows[i+1].cells[0]
        set_cell_bg(cell, bg)
        r = cell.paragraphs[0].add_run(s)
        r.font.size = Pt(9)
    doc.add_paragraph()


def kod_kutusu(doc, kod_satiri):
    """Gri arka planlı kod kutusu"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "F5F5F5")
    r = cell.paragraphs[0].add_run(kod_satiri)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    doc.add_paragraph()


def madde(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            r1 = p.add_run(item[0])
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = RGBColor(0x0D,0x47,0xA1)
            r2 = p.add_run(item[1])
            r2.font.size = Pt(10)
        else:
            r = p.add_run(item)
            r.font.size = Pt(10)


def build():
    doc = Document()
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # KAPAK
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("OPTİMİZASYON ÖNERİLERİ")
    r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = RGBColor(0x0D,0x47,0xA1)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("MQTT-SN vs CoAP — FIT IoT-LAB Projesi\nDeney Bulgularından Çıkan Somut İyileştirmeler")
    r2.font.size = Pt(13); r2.font.bold = True

    doc.add_paragraph()
    kutu(doc, "Hocanın sorusu ne anlama geliyor?", [
        "'Optimizasyon' sorusu şunu soruyor:",
        "  1. Deneyde ne yanlış gitti / neyi iyileştirebilirdin?",
        "  2. Gerçek dünyada bu protokolleri nasıl daha iyi çalıştırırsın?",
        "  3. Yaptığın ölçümlere bakarak hangi parametreleri değiştirirdin?",
        "Bizim PCap verilerinden 4 somut sorun çıktı → bunları düzeltmek = optimizasyon.",
    ], bg="FFF9C4", baslik_bg="F9A825")
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 1. DENEY BULGULARINDAKI SORUNLAR
    # ══════════════════════════════════════════════════
    baslik(doc, "1. Deneyde Ne Sorun Vardı? (Optimizasyon Nedenleri)")

    yaz(doc, "Optimizasyonu anlatmadan önce 'neden gerekli?' sorusunu cevaplamalısın. "
             "Bizim pcap verilerinden 4 somut sorun çıktı:", bold=True)
    doc.add_paragraph()

    sorunlar = [
        ("SORUN 1 — MQTT-SN Yıldız: Broker Aşırı Yüklenmesi", "C62828",
         ["mqttsn_star.pcap'te 8 kez CONNECT/CONNACK gözlemlendi.",
          "15 düğüm aynı anda broker'a bağlanmaya çalıştı → broker kesti.",
          "PINGRESP PDR'si %38.5'e düştü.",
          "SONUÇ: Yıldız topolojisinde MQTT-SN kullanılamaz hale geldi."],
         "FFEBEE"),

        ("SORUN 2 — CoAP Ağaç: %15.2 Paket Kaybı", "E65100",
         ["coap_tree.pcap: 33 CON gönderildi, 28 ACK alındı.",
          "5 paket 2 saniyelik timeout içinde ACK alamadı.",
          "Çok atlamalı ağaçta her atlama kayıp riski ekliyor.",
          "RPL ağacı henüz tam kararlı olmadan paket gönderilmiş olabilir."],
         "FFF3E0"),

        ("SORUN 3 — MQTT-SN Doğrusal/Yıldız: QoS-0 Kullanıldı", "1B5E20",
         ["mqttsn_linear ve mqttsn_star'da QoS-0 kullanıldı → PUBACK yok.",
          "Bu yüzden PDR ölçümü yapılamadı, PINGREQ'e bakıldı.",
          "QoS-0 = güvenilmez iletim. Gerçek uygulamada kabul edilemez.",
          "Ölçüm sınırlandı: sadece 2 veya 13 PINGREQ örneği vardı."],
         "E8F5E9"),

        ("SORUN 4 — RPL Yakınsama Süresi Beklenmedi", "4A148C",
         ["Düğümler flash'landıktan hemen sonra paket gönderildi.",
          "RPL ağacının tam kurulması 10-30 saniye alıyor.",
          "Bu sürede gönderilen paketler kayboldu veya yanlış yönlendi.",
          "CoAP ağaç başındaki kayıplar bu yüzden olmuş olabilir."],
         "F3E5F5"),
    ]

    for baslik_text, renk, satirlar, bg in sorunlar:
        kutu(doc, baslik_text, satirlar, bg=bg, baslik_bg=renk)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 2. PROTOKOL SEVİYESİ OPTİMİZASYONLAR
    # ══════════════════════════════════════════════════
    baslik(doc, "2. Protokol Seviyesi Optimizasyonlar")

    # 2.1 MQTT-SN
    baslik(doc, "2.1 MQTT-SN Optimizasyonları", level=2)

    yaz(doc, "A) QoS-1'i Tüm Topolojilerde Kullan", bold=True, renk="1565C0")
    yaz(doc, "Sorun: Yıldız ve doğrusal topolojilerde QoS-0 kullandık, PDR ölçülemedi.\n"
             "Çözüm: Her topolojide QoS-1 kullan → her PUBLISH için PUBACK beklenir → "
             "gerçek PDR ölçülür, güvenilirlik artar.")
    kutu(doc, "C kodunda değişiklik (mqtt_node/main.c):", [
        "// ÖNCE (QoS-0):",
        "emcute_pub(&pub_topic, payload, len, EMCUTE_QOS_0);",
        "",
        "// SONRA (QoS-1 — tüm topolojilerde):",
        "emcute_pub(&pub_topic, payload, len, EMCUTE_QOS_1);",
        "",
        "→ Her pakette PUBACK alınır → %100 PDR elde edilebilir.",
    ], bg="F5F5F5", baslik_bg="424242")

    yaz(doc, "B) Bağlantı Fırtınasını Önle — Exponential Backoff", bold=True, renk="1565C0")
    yaz(doc, "Sorun: 15 düğüm aynı anda bağlanmaya çalıştı → broker çöktü.\n"
             "Çözüm: Her düğüm rastgele bekleme süresiyle bağlansın.")
    kutu(doc, "C kodunda değişiklik:", [
        "// Her düğüm kendi node ID'sine göre farklı beklesin:",
        "uint32_t node_id = /* AT86RF231 adresinden al */;",
        "uint32_t backoff_ms = (node_id % 10) * 500;  // 0-5 saniye arası",
        "ztimer_sleep(ZTIMER_MSEC, backoff_ms);",
        "emcute_con(&gw, true, NULL, NULL, 0, 0);",
        "",
        "→ Düğümler 0.5 saniye arayla bağlanır → broker aşırı yüklenmez.",
    ], bg="F5F5F5", baslik_bg="424242")

    yaz(doc, "C) Keep-Alive Süresini Optimize Et", bold=True, renk="1565C0")
    yaz(doc, "Sorun: PINGREQ/PINGRESP trafiği gereksiz bant genişliği tüketir.\n"
             "Çözüm: Veri gönderim sıklığına göre keep-alive ayarla.")
    kutu(doc, "Makefile'da değişiklik:", [
        "# Varsayılan keep-alive = 10 saniye",
        "# Sensör 500ms'de bir veri gönderiyorsa keep-alive 60s olabilir:",
        "CFLAGS += -DCONFIG_EMCUTE_KEEPALIVE_PING=60",
        "",
        "→ PINGREQ trafiği azalır → ağ trafiği %20-30 düşer.",
    ], bg="F5F5F5", baslik_bg="424242")

    yaz(doc, "D) Hiyerarşik Broker (İki Kademeli)", bold=True, renk="1565C0")
    yaz(doc, "Sorun: Tek broker tüm trafiği taşıyor → darboğaz.\n"
             "Çözüm: Her ağaç kademesinde bir yerel broker, bunlar üst broker'a bağlı.")
    kutu(doc, "Mimari değişiklik:", [
        "Mevcut: 15 düğüm → 1 broker",
        "",
        "Optimizasyon:",
        "  Kademe-1 düğümler → Yerel Broker-1",
        "  Kademe-2 düğümler → Yerel Broker-2",
        "  Broker-1 + Broker-2 → Ana Broker",
        "",
        "→ Her broker maksimum 5 düğüm taşır → aşırı yüklenme önlenir.",
    ], bg="E3F2FD", baslik_bg="1565C0")

    doc.add_paragraph()

    # 2.2 CoAP
    baslik(doc, "2.2 CoAP Optimizasyonları", level=2)

    yaz(doc, "A) CON Timeout ve Tekrar Deneme Ayarı", bold=True, renk="E65100")
    yaz(doc, "Sorun: coap_tree'de 5 paket ACK alamadı, 2 saniyede timeout oldu.\n"
             "Çözüm: Çok atlamalı ağaçta gecikme daha fazla olabilir — timeout'u artır.")
    kutu(doc, "RIOT OS CoAP parametresi (Makefile veya main.c):", [
        "# Varsayılan ACK_TIMEOUT = 2 saniye",
        "# 3 kademeli ağaç için 4 saniye daha uygun:",
        "CFLAGS += -DCONFIG_GCOAP_RESP_WAIT_MS=4000",
        "",
        "# Maksimum tekrar deneme sayısı (varsayılan 4):",
        "CFLAGS += -DCONFIG_COAP_MAX_RETRANSMIT=5",
        "",
        "→ Ağaçtaki gecikmelere tolerans artar → PDR %84.8'den %95+'a çıkabilir.",
    ], bg="F5F5F5", baslik_bg="424242")

    yaz(doc, "B) CoAP Observe — Sürekli Polling Yerine Push", bold=True, renk="E65100")
    yaz(doc, "Sorun: Her istemci her 500ms'de bir CON request gönderiyor → yüksek trafik.\n"
             "Çözüm: CoAP Observe (RFC 7641) — istemci bir kez subscribe eder, "
             "sunucu değişince kendisi gönderir.")
    kutu(doc, "Mimari değişiklik:", [
        "Mevcut: İstemci her 500ms POST gönder → sunucu yanıtla",
        "",
        "Observe ile:",
        "  1. İstemci bir kez GET /metrics/rtt Observe:0 gönderir",
        "  2. Sunucu değer değiştiğinde otomatik bildirim gönderir",
        "  3. İstemci tekrar istemek zorunda kalmaz",
        "",
        "→ Ağ trafiği dramatik azalır, özellikle çok atlamalı topolojilerde.",
    ], bg="FFF3E0", baslik_bg="E65100")

    yaz(doc, "C) NON + Uygulama Katmanı Onayı", bold=True, renk="E65100")
    yaz(doc, "Sorun: CON/ACK her pakette ekstra gecikme ve trafik.\n"
             "Çözüm: NON mesaj kullan, uygulama katmanında sıra numarasıyla PDR hesapla.")
    kutu(doc, "Yaklaşım:", [
        "// NON gönder (daha hızlı):",
        "gcoap_req_init(&pdu, buf, len, COAP_METHOD_POST, \"/metrics/rtt\");",
        "// CON yerine NON kullan → ACK beklenmez",
        "",
        "// Uygulama katmanında:",
        "// Payload'a seq numarası göm → alıcıda eksik seq = kayıp paketi tespit et",
        "snprintf(payload, sizeof(payload), \"%u|%lu\", seq, timestamp);",
        "",
        "→ RTT ölçümü seq-bazlı yapılır, CON overhead ortadan kalkar.",
    ], bg="F5F5F5", baslik_bg="424242")

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 3. AĞ SEVİYESİ OPTİMİZASYONLAR
    # ══════════════════════════════════════════════════
    baslik(doc, "3. Ağ Seviyesi Optimizasyonlar (RPL / 6LoWPAN)")

    baslik(doc, "3.1 RPL Yakınsama Bekleme Süresi", level=2)
    yaz(doc, "Sorun: Düğümler flash'landıktan hemen sonra paket gönderdi → "
             "RPL henüz tamamlanmamış → erken paketler kayboldu.\n"
             "Çözüm: Yakınsama için yeterli süre bekle.")
    kutu(doc, "C kodunda düzeltme:", [
        "// Şu an: 5 saniye bekleniyor",
        "#define WARMUP_MS  (5000U)",
        "",
        "// Optimize: RPL yakınsama için 15-20 saniye bekle",
        "#define WARMUP_MS  (20000U)",
        "",
        "// Daha iyi: RPL'in hazır olduğunu kontrol et",
        "// Shell'de: rpl komutunu çalıştır, DODAG kökü görünene kadar bekle",
        "",
        "→ İlk 5-10 paketteki kayıplar önlenir.",
    ], bg="F5F5F5", baslik_bg="424242")

    baslik(doc, "3.2 RPL Objective Function: OF0 yerine MRHOF", level=2)
    yaz(doc, "Sorun: Varsayılan RPL OF0 sadece hop sayısına bakıyor — zayıf bağlantıları da seçiyor.\n"
             "Çözüm: MRHOF (Minimum Rank with Hysteresis) + ETX metriği kullan.")
    kutu(doc, "RIOT OS Makefile değişikliği:", [
        "# OF0 (varsayılan — sadece hop sayısı):",
        "# USEMODULE += gnrc_rpl",
        "",
        "# MRHOF (ETX — beklenen iletim sayısı bazlı):",
        "USEMODULE += gnrc_rpl",
        "CFLAGS += -DCONFIG_GNRC_RPL_DEFAULT_OCP=1  # OCP=1 = MRHOF",
        "",
        "→ RPL daha güçlü bağlantıları seçer → PDR artar, RTT azalır.",
    ], bg="F5F5F5", baslik_bg="424242")
    yaz(doc, "ETX nedir? → Expected Transmissions. Bir paketi iletmek için ortalama kaç kez "
             "denemek gerektiğini ölçer. ETX=1 mükemmel, ETX=3 kötü bağlantı.", italic=True)

    baslik(doc, "3.3 TX Güç Optimizasyonu", level=2)
    yaz(doc, "Sorun: Sabit TX gücü — bazı bağlantılar çok güçlü (bant genişliği israfı), "
             "bazıları çok zayıf (kayıp).\n"
             "Çözüm: Bağlantı kalitesine göre dinamik TX gücü.")
    kutu(doc, "RIOT OS shell komutu veya kod:", [
        "// Bağlantı kalitesini ölç:",
        "ifconfig 7  // RSSI değerini gör",
        "",
        "// TX gücünü ayarla:",
        "ifconfig 7 set txpower -7  // dBm cinsinden",
        "",
        "// Dinamik ayar için:",
        "// RSSI > -60 dBm → txpower düşür (gereksiz güç harcama)",
        "// RSSI < -85 dBm → txpower artır (bağlantı zayıf)",
        "",
        "→ Enerji tüketimi azalır, ağ ömrü uzar.",
    ], bg="E8F5E9", baslik_bg="2E7D32")

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 4. ENERJİ OPTİMİZASYONU
    # ══════════════════════════════════════════════════
    baslik(doc, "4. Enerji Optimizasyonu (Hoca Bunu da Sorabilir)")

    yaz(doc, "IoT'nin en kritik kısıtı: pil ömrü. Sensörler yıllarca çalışmalı.", bold=True)
    doc.add_paragraph()

    baslik(doc, "4.1 Duty Cycling — Uyku/Uyanık Döngüsü", level=2)
    yaz(doc, "Düğüm radyosunu sürekli açık tutmak yerine, paket gönder → uyu → uyan → gönder.")
    kutu(doc, "C kodu değişikliği:", [
        "// Mevcut: sürekli aktif",
        "for (i = 0; i < NUM_MESSAGES; i++) {",
        "    emcute_pub(...);",
        "    ztimer_sleep(ZTIMER_MSEC, 500);  // sadece 500ms bekle",
        "}",
        "",
        "// Enerji optimizasyonu: 500ms veri gönder, 9500ms uyu",
        "// %5 duty cycle → %95 enerji tasarrufu",
        "for (i = 0; i < NUM_MESSAGES; i++) {",
        "    emcute_pub(...);",
        "    pm_set(1);  // RIOT low-power mode",
        "    ztimer_sleep(ZTIMER_MSEC, 9500);",
        "    pm_unblock(1);",
        "}",
    ], bg="F5F5F5", baslik_bg="424242")

    baslik(doc, "4.2 Mesaj Birleştirme (Aggregation)", level=2)
    yaz(doc, "Sorun: Her sensör okuma için ayrı paket → çok fazla radyo açılış/kapanış.\n"
             "Çözüm: 10 ölçümü birleştir → tek pakette gönder.")
    kutu(doc, "Etki:", [
        "Mevcut: 10 paket × (radyo açılış + iletim + kapanış) = 10 × enerji",
        "Sonra:  1 paket × (radyo açılış + iletim + kapanış) = 1 × enerji",
        "",
        "→ Radyo açılış enerjisi en pahalı kısım → %60-70 enerji tasarrufu mümkün.",
        "→ Gecikme artar (trade-off: enerji vs. latency)",
    ], bg="E3F2FD", baslik_bg="1565C0")

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 5. ÖLÇÜM OPTİMİZASYONU
    # ══════════════════════════════════════════════════
    baslik(doc, "5. Deneyin Kendisini Nasıl İyileştirirdin")

    yaz(doc, "Hoca şunu da sorabilir: 'Daha iyi sonuç almak için ne değiştirirdin?'", bold=True, renk="C62828")
    doc.add_paragraph()

    iyilestirmeler = [
        ("Daha Fazla Paket Gönder",
         ["Mevcut: 11 MQTT-SN paketi (çok az, istatistiksel güven düşük)",
          "İyileştirme: 500-1000 paket gönder → standart sapma ve güven aralığı hesaplanabilir",
          "Deney süresini 120 dakikaya çıkar → yeterli örnek toplanır"]),

        ("Birden Fazla Çalıştır (Multiple Runs)",
         ["Her topoloji için tek deney yaptık",
          "İyileştirme: Her topolojiyi 5 kez çalıştır → ortalamasını al",
          "Neden: Radyo ortamı değişken, tek çalıştırma yanıltıcı olabilir"]),

        ("Hop Sayısını Kaydet",
         ["Mevcut: Paket kaç atlama yaptığını bilmiyoruz",
          "İyileştirme: RPL hop count'u payload'a göm",
          "→ RTT vs hop sayısı grafiği çiz → çok atlamanın etkisi görülür"]),

        ("Ağaç Yapısını Doğrula",
         ["Mevcut: TX güçle topoloji zorladık ama gerçekten oluştu mu?",
          "İyileştirme: Deneyd önce her düğümde 'rpl' komutu çalıştır",
          "Parent'ları kaydet → gerçek ağaç yapısını belgele"]),

        ("Karşılaştırmalı Enerji Ölçümü",
         ["Mevcut: Enerji ölçümü yok",
          "İyileştirme: FIT IoT-LAB'ın built-in güç ölçüm özelliğini kullan",
          "iotlab-experiment submit -l 'a8-103,Power' → mA cinsinden ölçüm",
          "→ MQTT-SN vs CoAP enerji tüketimi karşılaştırması yapılabilir"]),
    ]

    for baslik_text, satirlar in iyilestirmeler:
        kutu(doc, baslik_text, satirlar, bg="E8EAF6", baslik_bg="283593")

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 6. HOCANIN SORUSUNA HAZIR CEVAPLAR
    # ══════════════════════════════════════════════════
    baslik(doc, "6. Hocanın Sorularına Hazır Cevaplar")

    sorular = [
        ("'Bu projeyi optimize etsen ne yapardın?'",
         "3 şeyi değiştirirdim:\n"
         "1) MQTT-SN'de tüm topolojilerde QoS-1 kullanırdım — tutarlı PDR ölçümü için.\n"
         "2) CoAP'ta CON timeout'u 4 saniyeye çıkarırdım — ağaç topolojisindeki %15 kaybı önlemek için.\n"
         "3) RPL için 20 saniye yakınsama beklerdim — erken paket kayıplarını önlemek için."),

        ("'RPL optimizasyonu nasıl yapılır?'",
         "Varsayılan OF0 objective function yerine MRHOF kullanırdım. "
         "OF0 sadece hop sayısına bakıyor, MRHOF ETX metriği kullanıyor — "
         "bağlantı kalitesini ölçerek daha iyi yol seçiyor. "
         "RIOT OS'ta CFLAGS += -DCONFIG_GNRC_RPL_DEFAULT_OCP=1 ile aktif edilir."),

        ("'Enerji optimizasyonu yapabilir miydin?'",
         "Evet. Duty cycling uygulardım — düğüm paket gönderip %95 süre uyur. "
         "Ayrıca mesaj aggregation yapardım: 10 ölçümü tek pakette göndermek "
         "radyo açılış/kapanış enerjisini 10 kattan 1 kata düşürür. "
         "FIT IoT-LAB'ın güç ölçüm özelliğiyle mA bazlı karşılaştırma da yapılabilirdi."),

        ("'Broker darboğazını nasıl çözerdin?'",
         "İki yol: Ya exponential backoff ekle (düğümler sırayla bağlansın), "
         "ya da hiyerarşik broker kullan (her kademe için ayrı RSMB, "
         "bunlar ana broker'a bağlı). "
         "İkincisi gerçek endüstriyel sistemlerde de kullanılan yaklaşım."),

        ("'CoAP Observe ne işe yarar?'",
         "Şu an istemci her 500ms'de CON POST gönderiyor — çok fazla trafik. "
         "CoAP Observe ile istemci bir kez abone olur, "
         "sunucu değer değişince otomatik bildirim gönderir. "
         "Bu özellikle çok atlamalı ağaçta trafik yükünü dramatik azaltır."),
    ]

    for soru, cevap in sorular:
        p = doc.add_paragraph()
        r1 = p.add_run("SORU: " + soru)
        r1.bold = True; r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0xC6,0x28,0x28)
        p2 = doc.add_paragraph()
        r2 = p2.add_run("CEVAP: " + cevap)
        r2.font.size = Pt(10); r2.italic = True
        p2.paragraph_format.left_indent = Cm(0.8)
        p2.paragraph_format.space_after = Pt(10)

    doc.add_page_break()

    # HIZLI ÖZET
    baslik(doc, "7. Hızlı Özet — Optimizasyon Başlıkları")
    kutu(doc, "Bunları ezberle, yeterlі:", [
        "1. QoS-1 her yerde → güvenilir iletim, ölçülebilir PDR",
        "2. Exponential backoff → broker aşırı yüklenme önlenir",
        "3. CoAP timeout artır → ağaç topolojisinde kayıp azalır",
        "4. RPL MRHOF + ETX → zayıf bağlantılardan kaçınır, daha iyi yol seçimi",
        "5. 20s yakınsama bekleme → erken paket kayıpları önlenir",
        "6. Duty cycling → %95 enerji tasarrufu (radyo uyku modu)",
        "7. Mesaj aggregation → 10 ölçüm = 1 paket → enerji tasarrufu",
        "8. CoAP Observe → polling yerine push → trafik azalır",
        "9. Hiyerarşik broker → tek broker darboğazı ortadan kalkar",
        "10. Çoklu çalıştırma (5 run) → istatistiksel güvenilirlik artar",
    ], bg="FFF9C4", baslik_bg="F57F17")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Bu optimizasyonların hepsi RIOT OS'ta uygulanabilir. Başarılar!")
    r.font.size = Pt(12); r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)

    doc.save(str(OUT))
    print(f"Optimizasyon dosyasi olusturuldu -> {OUT}")


if __name__ == "__main__":
    build()
